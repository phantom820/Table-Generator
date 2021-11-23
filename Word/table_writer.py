import json
from typing import List
import docx
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Cm
import numpy as np
import uuid
import os

# Word Table Writer
''' - Top & bottom rows bordered: Medium List
- Header and bottom bordered: Light Shading
- Partially bordered: Colorful Grid
- Fully bordered: Light Grid 
'''
class TableWriter:    
    def __init__(self):    
        with open('templates/paragraphs.json','r') as f:
            paragraphs = json.load(f)
        self.paragraphs = []
        for key in paragraphs:
            text = "".join(paragraphs[key])
            self.paragraphs.append(text)
       
    ''' write table '''
    def write_table(self,table,df):
        for j in range(df.shape[-1]):
            table.cell(0,j).text = df.columns[j]
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i+1,j).text = str(df.values[i,j])
                
    'make row of cells background colored, defaults to column header row'
    def color_table(self,table,color='1F5C8B'):
        color = "".join([r'<w:shd {} ','w:fill=',f'"{color}"','/>'])
        for row in table.rows:
            for cell in row.cells:
                shading_elm_2 = parse_xml(color.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_2)
                
    ''' structure margins'''
    def structure_margins(self,doc,vertical=0.5,horizontal=1):
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(vertical)
            section.bottom_margin = Cm(vertical)
            section.left_margin = Cm(horizontal)
            section.right_margin = Cm(horizontal)

    ''' just write a single table '''
    def add_table(self,doc,table_type:int,df,df_inner=None):
        p = np.random.uniform(0,1)
        if p>0.5:
            accent = ' Accent '+str(int(np.random.uniform(1,7)))
        elif p<0.5:
            accent = ''
        if table_type==0:
            table = doc.add_table(df.shape[0]+1, df.shape[1],style='Medium List 1'+accent)
            self.write_table(table,df)
            return doc
            
        elif table_type==1:
            table = doc.add_table(df.shape[0]+1, df.shape[1],'Light Shading'+accent)
            self.write_table(table,df)
            return doc
        
        elif table_type==2:
            table = doc.add_table(df.shape[0]+1, df.shape[1],'Colorful Grid'+accent)
            self.write_table(table,df)
            return doc
        
        elif table_type==3:
            table = doc.add_table(df.shape[0]+1, df.shape[1],'Light Grid'+accent)
            self.write_table(table,df)
            return doc
            
        elif  table_type==4:
            styles = ['Medium List 1','Light Shading','Colorful Grid','Light Grid']
            style = styles[int(np.random.uniform(0,len(styles)))]
            style = style+accent
            x = int(np.random.uniform(0,df.shape[-1]))
            y = int(np.random.uniform(1,df.shape[-1]))
            table = doc.add_table(df.shape[0]+1, df.shape[1],style='Light Grid'+accent)
            for j in range(df.shape[-1]):
                table.cell(0,j).text = df.columns[j]
            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    if i==x and j==y:
                        inner_table = table.cell(i+1,j).add_table(df_inner.shape[0]+1,df_inner.shape[1])
                        self.write_table(inner_table,df_inner)
                        inner_table.style = style
                    else: 
                        table.cell(i+1,j).text = str(df.values[i,j])
            return doc
             
    ''' write multiple tables '''
    def write(self,table_types:List[int],dfs,inner_dfs=None)->str:
        if inner_dfs is None:
            inner_dfs = [None for i in range(len(table_types))]
        doc = docx.Document()
        self.structure_margins(doc)
        for i in range(len(table_types)):
            text = self.paragraphs[i]
            doc.add_paragraph("".join(["\n",text]))
            header = doc.add_heading('Table '+str(i+1))
            header.alignment = 1
            df = dfs[i]
            if inner_dfs[i] is None:
                doc = self.add_table(doc,table_types[i],dfs[i])
            else:
                doc = self.add_table(doc,table_types[i],dfs[i],inner_dfs[i])
        doc.add_paragraph("".join(["\n\n",self.paragraphs[-1]]))
        _id = str(uuid.uuid4().int & (1<<64)-1)
        doc.save(f'temp/{_id}.docx')
        out_doc = docx.Document(f'temp/{_id}.docx')
        os.remove(f'temp/{_id}.docx')
        tables = out_doc.tables
        for table in tables:
            self.color_table(table)
        return doc,out_doc